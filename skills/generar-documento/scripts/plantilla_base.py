"""
PLANTILLA BASE — Documentos Profesionales (CAHZ)
Copiar, adaptar el contenido y ejecutar desde ~/Downloads/gen_doc_temp.py
Borrar el script temporal después de generar el .docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ── CONFIGURACIÓN ──────────────────────────────────────────────────────────────
FUENTE       = "Arial Narrow"
TAMANIO      = 12
COLOR_NEGRO  = RGBColor(0, 0, 0)
COLOR_GRIS   = RGBColor(0xD9, 0xD9, 0xD9)   # Encabezados de tabla neutro

# ── HELPERS ───────────────────────────────────────────────────────────────────

def configurar_margenes(doc, sup=3, inf=2.5, izq=3, der=2.5):
    for s in doc.sections:
        s.top_margin    = Cm(sup)
        s.bottom_margin = Cm(inf)
        s.left_margin   = Cm(izq)
        s.right_margin  = Cm(der)

def set_font(run, bold=False, italic=False, size=None, underline=False):
    run.font.name      = FUENTE
    run.font.size      = Pt(size or TAMANIO)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.underline = underline
    run.font.color.rgb = COLOR_NEGRO

def _base_para(doc, align, sp_before, sp_after):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(sp_before)
    pf.space_after  = Pt(sp_after)
    pf.line_spacing = Pt(15)
    return p

def add_para(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             bold=False, italic=False, size=None,
             sp_before=0, sp_after=6):
    p = _base_para(doc, align, sp_before, sp_after)
    if text:
        r = p.add_run(text)
        set_font(r, bold=bold, italic=italic, size=size)
    return p

def add_mixed(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              sp_before=0, sp_after=6):
    """parts = [(texto, bold, italic), ...]"""
    p = _base_para(doc, align, sp_before, sp_after)
    for text, bold, italic in parts:
        r = p.add_run(text)
        set_font(r, bold=bold, italic=italic)
    return p

def add_heading(doc, titulo, sp_before=10, sp_after=4):
    """Título de sección: Bold + Underline"""
    p = _base_para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY, sp_before, sp_after)
    r = p.add_run(titulo)
    set_font(r, bold=True, underline=True)
    return p

def add_table_row(table, row_idx, values, header=False):
    """Agrega valores a una fila y aplica estilo."""
    row = table.rows[row_idx]
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(val))
        set_font(run, bold=header)
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)
        if header:
            from docx.oxml import OxmlElement
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'D9D9D9')
            tcPr.append(shd)

# ── DOCUMENTO ─────────────────────────────────────────────────────────────────
doc = Document()
configurar_margenes(doc)

# 1. FECHA (derecha)
add_para(doc, "Bogotá D.C., [DÍA] de [MES] de [AÑO]",
         align=WD_ALIGN_PARAGRAPH.RIGHT, sp_after=14)

# 2. DESTINATARIO (izquierda)
add_para(doc, "Señor(a)", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, sp_after=2)
add_para(doc, "[NOMBRE COMPLETO DESTINATARIO]",
         align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, sp_after=2)
add_para(doc, "[Cargo]", align=WD_ALIGN_PARAGRAPH.LEFT, sp_after=2)
add_para(doc, "[Entidad]", align=WD_ALIGN_PARAGRAPH.LEFT, sp_after=2)
add_para(doc, "[Ciudad]", align=WD_ALIGN_PARAGRAPH.LEFT, sp_after=14)

# 3. ASUNTO / RADICADO
add_mixed(doc, [
    ("Asunto:\t", True, False),
    ("[Descripción concisa del asunto]", False, False),
], sp_after=4)
add_mixed(doc, [
    ("Radicado:\t", True, False),
    ("[No. de radicado]", False, False),
], sp_after=14)

# 4. SALUDO
add_para(doc, "Respetado(a) señor(a):", align=WD_ALIGN_PARAGRAPH.LEFT, sp_after=8)

# 5. PÁRRAFO INTRODUCTORIO
add_para(doc,
    "[Párrafo de apertura que contextualiza la respuesta, cita el acto o norma "
    "que motiva el documento y anuncia la estructura de la respuesta.]",
    sp_after=10)

# 6. CUERPO — ejemplo con dos secciones
add_heading(doc, "Frente al punto 1: [Título del punto]")
add_para(doc, "[Desarrollo del punto 1. Redacción directa, verbos activos, "
              "máximo 4 líneas por párrafo.]")

add_heading(doc, "Frente al punto 2: [Título del punto]")
add_para(doc, "[Desarrollo del punto 2.]")

# 7. CIERRE
add_para(doc, "", sp_after=6)
add_para(doc,
    "[Párrafo de cierre. Reitera el compromiso institucional o indica el "
    "siguiente paso. Citar norma de confidencialidad si aplica.]",
    sp_before=6, sp_after=20)

# 8. FIRMA
add_para(doc, "Atentamente,", align=WD_ALIGN_PARAGRAPH.LEFT, sp_after=60)
add_para(doc, "[NOMBRE COMPLETO EN MAYÚSCULAS]",
         align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, sp_after=2)
add_para(doc, "[Cargo]", align=WD_ALIGN_PARAGRAPH.LEFT, sp_after=2)
add_para(doc, "[Entidad — Sigla]", align=WD_ALIGN_PARAGRAPH.LEFT, sp_after=20)

# 9. ELABORÓ / REVISÓ
add_mixed(doc, [("Proyecto: ", True, False), ("[Nombre] — [Cargo]", False, False)],
          align=WD_ALIGN_PARAGRAPH.LEFT, sp_after=2)
add_mixed(doc, [("Revisó:   ", True, False), ("[Nombre] — [Cargo]", False, False)],
          align=WD_ALIGN_PARAGRAPH.LEFT, sp_after=0)

# ── GUARDAR ───────────────────────────────────────────────────────────────────
import os
ruta = os.path.expanduser("~/Downloads/DOCUMENTO_PROFESIONAL_v1.docx")
doc.save(ruta)
print(f"✓ Guardado: {ruta}")
